from docx import Document
import os
import uuid
import sys

def get_base_path():
    """uploads和outputs在exe同级目录"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

BASE_PATH = get_base_path()

UPLOADS_DIR = os.path.join(BASE_PATH, "uploads")
OUTPUTS_DIR = os.path.join(BASE_PATH, "outputs")


def translate_docx(input_path: str, translate_fn) -> str:

    """
    解析docx文件，逐段翻译，生成译文文件
    返回译文文件路径
    """
    doc = Document(input_path)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated = translate_fn(paragraph.text)
            paragraph.text = translated

    # 处理表格内文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        translated = translate_fn(paragraph.text)
                        paragraph.text = translated

    # 生成输出文件
    output_filename = f"translated_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(OUTPUTS_DIR, output_filename)
    doc.save(output_path)

    return output_path



def translate_txt(input_path: str, translate_fn) -> str:
    """
    解析txt文件，逐段翻译，生成译文文件
    返回译文文件路径
    """
    # 读取原文
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按段落分割（按空行分割）
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    
    # 翻译每个段落
    translated_paragraphs = []
    for para in paragraphs:
        if para.strip():
            translated = translate_fn(para)
            translated_paragraphs.append(translated)
        else:
            translated_paragraphs.append('')
    
    # 重新组合
    translated_content = '\n\n'.join(translated_paragraphs)
    
    # 生成输出文件
    output_filename = f"translated_{uuid.uuid4().hex[:8]}.txt"
    output_path = os.path.join(OUTPUTS_DIR, output_filename)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(translated_content)
    
    return output_path